# -*- coding: utf-8 -*-
# @Time    : 2023/6/2 8:30
# @Author  : Monarch
# @File    : word_table_extract.py
# @Software: PyCharm

import docx
import os
import pandas as pd
import re
from glob import glob

os.chdir(r'D:\Work\需解决')
data = []
reg = '<w:t>(.+?)</w:t>'
cell_null = 'null'
old_columns = ['工作单位', '姓名', '职务/职称', '是否报告', '论文题目']
columns = ['单位', '姓名', '教师或学生', '是否报告', '报告题目']

files = glob('11/*.docx')
# file = r'回执1/河南财经政法大学资源与环境学院 罗庆教授--参会回执 (1).docx'
for file in files:
    doc = docx.Document(file)
    # 按标题判断是否为回执
    title = ''
    for i in doc.paragraphs:
        title = i.text.split()
        if title != '':
            break
    if ('回执' not in title) and ('注册表' not in title) and ('学术年会' not in title):
        continue

    table = doc.tables[0]
    texts = []

    for row in table.rows:
        for cell in row.cells:
            if cell.text != cell_null:              # 去重
                # 单元格转化为xml字符串
                cell_xml = cell._element.xml
                # 规范符号
                cell_xml = cell_xml.replace('■', '√')
                # 特殊符号转化为√
                cell_xml = cell_xml.replace('w:char="F052"', '<w:t>√</w:t>')
                cell_xml = cell_xml.replace('w:char="0052"', '<w:t>√</w:t>')
                cell_xml = cell_xml.replace('w:char="F0FE"', '<w:t>√</w:t>')
                # 正则提取
                cell_text_list = re.findall(reg, cell_xml, re.S)
                # 合并
                cell_text = ''.join(cell_text_list)
                cell.text = cell_null               # 去重
                texts.append(cell_text)
    # 转为dict
    table_dict = {}
    for i in range(0, len(texts), 2):
        table_dict[texts[i]] = texts[i+1]
    # 判断是否作报告
    if (table_dict['是否报告'] == '√不作报告作报告') or (table_dict['是否报告'] == '') or (table_dict['论文题目'] == ''):
        table_dict['是否报告'] = '否'
    else:
        table_dict['是否报告'] = '是'
    data.append(table_dict)
    break


# 转为DataFrame
# df = pd.DataFrame(data)
# # 提取需要的列
# save_df = df[old_columns]
# # 修改列名
# save_df.columns = columns
# # 序号列
# save_df.index.name = '序号'
# save_df.index = save_df.index+1
# # 保存
# save_df.to_excel('回执11.xlsx')


